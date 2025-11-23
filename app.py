import os
import io
from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import PyPDF2
import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY
import shutil

from api_manager import APIKeyManager
from ai_providers import get_provider

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'

# Create necessary folders
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# Initialize API Key Manager
api_manager = APIKeyManager()

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        # Try pdfplumber first (better for complex PDFs)
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed, trying PyPDF2: {e}")
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e2:
            print(f"PyPDF2 also failed: {e2}")
            raise Exception("Could not extract text from PDF")

    return text.strip()

def extract_text_from_docx(file_path):
    """Extract text from Word document"""
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return '\n'.join(text)

def extract_text_from_txt(file_path):
    """Extract text from text file"""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def extract_text(file_path, filename):
    """Extract text based on file type"""
    ext = filename.rsplit('.', 1)[1].lower()

    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    elif ext == 'txt':
        return extract_text_from_txt(file_path)
    else:
        raise Exception(f"Unsupported file type: {ext}")

def summarize_text(text):
    """Summarize text using configured AI provider"""
    # Get default provider
    provider_name = api_manager.get_default_provider()
    if not provider_name:
        raise Exception("No AI provider configured. Please configure an API key in settings.")

    # Get API key for provider
    api_key = api_manager.get_api_key(provider_name)
    if not api_key:
        raise Exception(f"API key not found for {provider_name}")

    # Get provider instance
    provider = get_provider(provider_name, api_key)
    if not provider:
        raise Exception(f"Unsupported provider: {provider_name}")

    # Summarize using the provider
    try:
        return provider.summarize(text)
    except Exception as e:
        raise Exception(f"Error summarizing text with {provider_name}: {str(e)}")

def create_pdf_output(summary, original_filename):
    """Create PDF file with summary"""
    output_filename = f"summary_{original_filename.rsplit('.', 1)[0]}.pdf"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

    # Create PDF
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)

    # Container for the 'Flowable' objects
    elements = []

    # Define styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY))

    # Add title
    title_style = styles['Heading1']
    title = Paragraph(f"Summary of {original_filename}", title_style)
    elements.append(title)
    elements.append(Spacer(1, 12))

    # Add summary text
    summary_style = styles['BodyText']
    summary_style.alignment = TA_JUSTIFY

    # Split summary into paragraphs
    for para in summary.split('\n'):
        if para.strip():
            p = Paragraph(para, summary_style)
            elements.append(p)
            elements.append(Spacer(1, 12))

    # Build PDF
    doc.build(elements)

    return output_path

def create_docx_output(summary, original_filename):
    """Create Word document with summary"""
    output_filename = f"summary_{original_filename.rsplit('.', 1)[0]}.docx"
    output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

    # Create Word document
    doc = Document()

    # Add title
    doc.add_heading(f"Summary of {original_filename}", 0)

    # Add summary text
    for para in summary.split('\n'):
        if para.strip():
            doc.add_paragraph(para)

    # Save document
    doc.save(output_path)

    return output_path

# Routes

@app.route('/')
def index():
    """Render main page or redirect to setup if not configured"""
    if not api_manager.has_any_provider():
        return redirect(url_for('setup'))
    return render_template('index.html')

@app.route('/setup')
def setup():
    """Render setup wizard"""
    return render_template('setup.html')

@app.route('/settings')
def settings():
    """Render settings page"""
    if not api_manager.has_any_provider():
        return redirect(url_for('setup'))
    return render_template('settings.html')

@app.route('/summarize', methods=['POST'])
def summarize():
    """Handle summarization request"""
    try:
        # Check if configured
        if not api_manager.has_any_provider():
            return jsonify({'error': 'No AI provider configured. Please complete setup first.'}), 400

        text_to_summarize = ""
        original_filename = "document"

        # Check if text was pasted/typed
        if 'manual_text' in request.form and request.form['manual_text'].strip():
            text_to_summarize = request.form['manual_text'].strip()
            original_filename = "manual_input.txt"

        # Check if file was uploaded
        elif 'file' in request.files:
            file = request.files['file']
            if file and file.filename and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                original_filename = filename
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)

                # Extract text from file
                text_to_summarize = extract_text(filepath, filename)

                # Clean up uploaded file
                os.remove(filepath)
            else:
                return jsonify({'error': 'Invalid file type. Please upload PDF, Word, or text file.'}), 400
        else:
            return jsonify({'error': 'No input provided. Please upload a file or paste text.'}), 400

        # Check if text is not empty
        if not text_to_summarize or len(text_to_summarize.strip()) < 50:
            return jsonify({'error': 'Text is too short to summarize. Please provide more content.'}), 400

        # Summarize the text
        summary = summarize_text(text_to_summarize)

        # Get output format preference
        output_format = request.form.get('output_format', 'pdf')

        # Generate output file
        if output_format == 'word':
            output_path = create_docx_output(summary, original_filename)
        else:
            output_path = create_pdf_output(summary, original_filename)

        # Return success with download info
        return jsonify({
            'success': True,
            'summary': summary,
            'download_filename': os.path.basename(output_path),
            'output_format': output_format
        })

    except Exception as e:
        print(f"Error in summarize: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download(filename):
    """Download generated summary file"""
    try:
        filepath = os.path.join(app.config['OUTPUT_FOLDER'], secure_filename(filename))
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# API Routes for Setup and Settings

@app.route('/api/test-key', methods=['POST'])
def test_api_key():
    """Test if an API key is valid"""
    try:
        data = request.json
        provider_name = data.get('provider')
        api_key = data.get('api_key')

        if not provider_name or not api_key:
            return jsonify({'success': False, 'error': 'Missing provider or API key'}), 400

        # Get provider instance
        provider = get_provider(provider_name, api_key)
        if not provider:
            return jsonify({'success': False, 'error': 'Invalid provider'}), 400

        # Test connection
        is_valid = provider.test_connection()

        if is_valid:
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': 'API key validation failed'}), 400

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/setup', methods=['POST'])
def save_setup():
    """Save initial setup configuration"""
    try:
        data = request.json
        provider = data.get('provider')
        api_key = data.get('api_key')

        if not provider or not api_key:
            return jsonify({'success': False, 'error': 'Missing provider or API key'}), 400

        # Add provider to config
        api_manager.add_provider(provider, api_key, set_as_default=True)

        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/providers', methods=['GET'])
def list_providers():
    """Get list of configured providers"""
    try:
        providers = api_manager.list_providers()
        return jsonify({'success': True, 'providers': providers})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/add-provider', methods=['POST'])
def add_provider():
    """Add a new provider"""
    try:
        data = request.json
        provider = data.get('provider')
        api_key = data.get('api_key')
        set_as_default = data.get('set_as_default', False)

        if not provider or not api_key:
            return jsonify({'success': False, 'error': 'Missing provider or API key'}), 400

        api_manager.add_provider(provider, api_key, set_as_default=set_as_default)

        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/remove-provider', methods=['POST'])
def remove_provider():
    """Remove a provider"""
    try:
        data = request.json
        provider = data.get('provider')

        if not provider:
            return jsonify({'success': False, 'error': 'Missing provider name'}), 400

        api_manager.remove_provider(provider)

        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/set-default', methods=['POST'])
def set_default_provider():
    """Set default provider"""
    try:
        data = request.json
        provider = data.get('provider')

        if not provider:
            return jsonify({'success': False, 'error': 'Missing provider name'}), 400

        api_manager.set_default_provider(provider)

        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/health')
def health():
    """Health check endpoint"""
    has_config = api_manager.has_any_provider()
    default_provider = api_manager.get_default_provider()

    return jsonify({
        'status': 'healthy',
        'configured': has_config,
        'default_provider': default_provider
    })

# Clean up old files on startup
def cleanup_folders():
    """Clean up upload and output folders"""
    for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
        if os.path.exists(folder):
            shutil.rmtree(folder)
        os.makedirs(folder, exist_ok=True)

if __name__ == '__main__':
    cleanup_folders()
    print("\n" + "="*60)
    print("📄 Document Summarizer is starting...")
    print("="*60)

    if not api_manager.has_any_provider():
        print("\n⚠️  No AI provider configured yet!")
        print("Please complete the setup wizard in your browser.")
        print("="*60 + "\n")
    else:
        default = api_manager.get_default_provider()
        print(f"\n✅ Using {default.upper()} as the default AI provider")
        print("="*60 + "\n")

    print("🌐 Open your browser and go to: http://localhost:5000")
    print("\n💡 Press Ctrl+C to stop the server\n")

    app.run(debug=False, host='0.0.0.0', port=5000)
