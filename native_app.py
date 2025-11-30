"""
Document Summarizer - Native Desktop Application
Professional multi-provider AI document summarization tool
Built with CustomTkinter for a modern, polished interface
"""

__version__ = "1.1.0"
__app_name__ = "Document Summarizer"

import os
import sys
import json
import threading
import webbrowser
import customtkinter as ctk
from tkinter import filedialog, messagebox
from tkinterdnd2 import DND_FILES, TkinterDnD
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY

from api_manager import APIKeyManager, get_config_dir
from ai_providers import get_provider, get_all_providers, get_provider_info, PROVIDERS

# Configure CustomTkinter
ctk.set_default_color_theme("blue")


class ThemeManager:
    """Manages application themes with persistence"""

    LIGHT_THEME = {
        "name": "light",
        "appearance_mode": "light",
        "BG_PRIMARY": "#FFFFFF",
        "BG_SECONDARY": "#F8FAFC",
        "BG_TERTIARY": "#F1F5F9",
        "TEXT_PRIMARY": "#0F172A",
        "TEXT_SECONDARY": "#475569",
        "TEXT_MUTED": "#94A3B8",
        "ACCENT_PRIMARY": "#2563EB",
        "ACCENT_HOVER": "#1D4ED8",
        "ACCENT_LIGHT": "#DBEAFE",
        "SUCCESS": "#059669",
        "SUCCESS_HOVER": "#047857",
        "SUCCESS_LIGHT": "#D1FAE5",
        "ERROR": "#EF4444",
        "WARNING": "#F59E0B",
        "BORDER": "#E2E8F0",
        "DIVIDER": "#CBD5E1",
        "TAB_ACTIVE": "#FFFFFF",
        "TAB_INACTIVE": "#E2E8F0",
        "TAB_HOVER": "#F1F5F9",
        "CARD_BG": "#FFFFFF",
        "INPUT_BG": "#F8FAFC",
    }

    DARK_THEME = {
        "name": "dark",
        "appearance_mode": "dark",
        "BG_PRIMARY": "#0F172A",
        "BG_SECONDARY": "#1E293B",
        "BG_TERTIARY": "#334155",
        "TEXT_PRIMARY": "#F8FAFC",
        "TEXT_SECONDARY": "#CBD5E1",
        "TEXT_MUTED": "#64748B",
        "ACCENT_PRIMARY": "#3B82F6",
        "ACCENT_HOVER": "#2563EB",
        "ACCENT_LIGHT": "#1E3A5F",
        "SUCCESS": "#10B981",
        "SUCCESS_HOVER": "#059669",
        "SUCCESS_LIGHT": "#064E3B",
        "ERROR": "#F87171",
        "WARNING": "#FBBF24",
        "BORDER": "#334155",
        "DIVIDER": "#475569",
        "TAB_ACTIVE": "#1E293B",
        "TAB_INACTIVE": "#0F172A",
        "TAB_HOVER": "#334155",
        "CARD_BG": "#1E293B",
        "INPUT_BG": "#334155",
    }

    PROVIDER_COLORS = {
        "gemini": "#4285F4",
        "openai": "#10A37F",
        "claude": "#CC785C",
        "groq": "#F55036",
        "grok": "#1DA1F2",
    }

    def __init__(self):
        self.config_path = os.path.join(get_config_dir(), "theme_config.json")
        self.current_theme = self.LIGHT_THEME
        self.load_theme()

    def load_theme(self):
        """Load theme from config file"""
        try:
            if os.path.exists(self.config_path):
                with open(self.config_path, 'r') as f:
                    config = json.load(f)
                    if config.get("theme") == "dark":
                        self.current_theme = self.DARK_THEME
                    else:
                        self.current_theme = self.LIGHT_THEME
        except Exception:
            self.current_theme = self.LIGHT_THEME

        ctk.set_appearance_mode(self.current_theme["appearance_mode"])

    def save_theme(self):
        """Save theme preference to config file"""
        try:
            with open(self.config_path, 'w') as f:
                json.dump({"theme": self.current_theme["name"]}, f)
        except Exception:
            pass

    def toggle_theme(self):
        """Toggle between light and dark themes"""
        if self.current_theme["name"] == "light":
            self.current_theme = self.DARK_THEME
        else:
            self.current_theme = self.LIGHT_THEME

        ctk.set_appearance_mode(self.current_theme["appearance_mode"])
        self.save_theme()
        return self.current_theme

    def is_dark(self) -> bool:
        """Check if current theme is dark"""
        return self.current_theme["name"] == "dark"

    def get(self, key: str) -> str:
        """Get a theme color value"""
        return self.current_theme.get(key, "#FFFFFF")


# Legacy Colors class for backward compatibility during transition
class Colors:
    # These will be updated dynamically based on theme
    BG_PRIMARY = "#FFFFFF"
    BG_SECONDARY = "#F8FAFC"
    BG_TERTIARY = "#F1F5F9"
    TEXT_PRIMARY = "#0F172A"
    TEXT_SECONDARY = "#475569"
    TEXT_MUTED = "#94A3B8"
    BLUE_PRIMARY = "#2563EB"
    BLUE_HOVER = "#1D4ED8"
    BLUE_LIGHT = "#DBEAFE"
    GREEN_PRIMARY = "#059669"
    GREEN_HOVER = "#047857"
    GREEN_LIGHT = "#D1FAE5"
    SUCCESS = "#10B981"
    ERROR = "#EF4444"
    WARNING = "#F59E0B"
    BORDER = "#E2E8F0"
    DIVIDER = "#CBD5E1"

    PROVIDER_COLORS = {
        "gemini": "#4285F4",
        "openai": "#10A37F",
        "claude": "#CC785C",
        "groq": "#F55036",
        "grok": "#1DA1F2",
    }


class DocumentSummarizerApp(TkinterDnD.Tk):
    """Main application window with professional tab-based UI"""

    def __init__(self):
        super().__init__()

        # Initialize theme manager first
        self.theme = ThemeManager()

        self.title("Document Summarizer")
        self.geometry("1000x700")
        self.minsize(850, 600)
        self.configure(bg=self.theme.get("BG_PRIMARY"))

        # Initialize API manager
        self.api_manager = APIKeyManager()

        # Track current state
        self.current_file_path: Optional[str] = None
        self.extracted_text: Optional[str] = None
        self.current_summary: Optional[str] = None
        self.original_filename: str = "document"
        self.current_tab: str = "summarize"

        # Create main container
        self.main_container = ctk.CTkFrame(self, fg_color=self.theme.get("BG_PRIMARY"), corner_radius=0)
        self.main_container.pack(fill="both", expand=True)

        # Check if setup is needed
        if not self.api_manager.has_any_provider():
            self.show_setup_view()
        else:
            self.show_main_view()

    def clear_container(self):
        """Clear all widgets from main container"""
        for widget in self.main_container.winfo_children():
            widget.destroy()

    def refresh_ui(self):
        """Refresh the entire UI with current theme"""
        self.configure(bg=self.theme.get("BG_PRIMARY"))
        self.main_container.configure(fg_color=self.theme.get("BG_PRIMARY"))
        if self.api_manager.has_any_provider():
            self.show_main_view()
        else:
            self.show_setup_view()

    # ===================== SETUP VIEW =====================

    def show_setup_view(self):
        """Show the setup wizard for first-time configuration"""
        self.clear_container()

        # Main setup container with max width
        setup_wrapper = ctk.CTkFrame(self.main_container, fg_color=Colors.BG_PRIMARY)
        setup_wrapper.pack(fill="both", expand=True)

        # Scrollable frame for setup content
        setup_frame = ctk.CTkFrame(setup_wrapper, fg_color=Colors.BG_PRIMARY, width=600)
        setup_frame.place(relx=0.5, rely=0.5, anchor="center")

        # Header
        header_frame = ctk.CTkFrame(setup_frame, fg_color="transparent")
        header_frame.pack(fill="x", pady=(0, 30))

        title_label = ctk.CTkLabel(
            header_frame,
            text="Document Summarizer",
            font=ctk.CTkFont(size=36, weight="bold"),
            text_color=Colors.TEXT_PRIMARY
        )
        title_label.pack()

        subtitle_label = ctk.CTkLabel(
            header_frame,
            text="Connect your AI provider to get started",
            font=ctk.CTkFont(size=16),
            text_color=Colors.TEXT_SECONDARY
        )
        subtitle_label.pack(pady=(8, 0))

        # Provider selection card
        provider_card = ctk.CTkFrame(setup_frame, fg_color=Colors.BG_SECONDARY, corner_radius=16)
        provider_card.pack(fill="x", pady=(0, 20))

        provider_header = ctk.CTkLabel(
            provider_card,
            text="Choose AI Provider",
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=Colors.TEXT_PRIMARY
        )
        provider_header.pack(pady=(24, 16), padx=24, anchor="w")

        # Provider buttons
        providers_frame = ctk.CTkFrame(provider_card, fg_color="transparent")
        providers_frame.pack(fill="x", padx=24, pady=(0, 24))

        self.selected_provider = ctk.StringVar(value="gemini")
        self.provider_buttons = {}

        all_providers = get_all_providers()
        row = 0
        col = 0

        for provider_name, info in all_providers.items():
            btn_frame = ctk.CTkFrame(providers_frame, fg_color="transparent")
            btn_frame.grid(row=row, column=col, padx=5, pady=5, sticky="ew")

            btn = ctk.CTkButton(
                btn_frame,
                text=info['display_name'],
                width=170,
                height=50,
                font=ctk.CTkFont(size=14),
                fg_color=Colors.BG_PRIMARY,
                text_color=Colors.TEXT_PRIMARY,
                border_width=2,
                border_color=Colors.BORDER,
                hover_color=Colors.BLUE_LIGHT,
                corner_radius=10,
                command=lambda p=provider_name: self.select_provider(p)
            )
            btn.pack(fill="x")
            self.provider_buttons[provider_name] = btn

            col += 1
            if col >= 3:
                col = 0
                row += 1

        # Configure grid columns
        for i in range(3):
            providers_frame.columnconfigure(i, weight=1)

        # API Key input card
        key_card = ctk.CTkFrame(setup_frame, fg_color=Colors.BG_SECONDARY, corner_radius=16)
        key_card.pack(fill="x", pady=(0, 20))

        key_header = ctk.CTkLabel(
            key_card,
            text="Enter API Key",
            font=ctk.CTkFont(size=18, weight="bold"),
            text_color=Colors.TEXT_PRIMARY
        )
        key_header.pack(pady=(24, 8), padx=24, anchor="w")

        self.api_help_label = ctk.CTkLabel(
            key_card,
            text="Get your free API key from Google AI Studio",
            font=ctk.CTkFont(size=13),
            text_color=Colors.TEXT_SECONDARY
        )
        self.api_help_label.pack(padx=24, anchor="w")

        # API key entry with show/hide
        key_entry_frame = ctk.CTkFrame(key_card, fg_color="transparent")
        key_entry_frame.pack(fill="x", padx=24, pady=(12, 8))

        self.api_key_entry = ctk.CTkEntry(
            key_entry_frame,
            placeholder_text="Paste your API key here...",
            height=50,
            font=ctk.CTkFont(size=14),
            show="*",
            border_width=2,
            border_color=Colors.BORDER,
            corner_radius=10
        )
        self.api_key_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))

        self.show_key_var = ctk.BooleanVar(value=False)
        show_key_btn = ctk.CTkButton(
            key_entry_frame,
            text="Show",
            width=70,
            height=50,
            font=ctk.CTkFont(size=13),
            fg_color=Colors.BG_TERTIARY,
            text_color=Colors.TEXT_SECONDARY,
            hover_color=Colors.BORDER,
            corner_radius=10,
            command=self.toggle_key_visibility
        )
        show_key_btn.pack(side="right")

        # Get API key link
        self.get_key_link = ctk.CTkButton(
            key_card,
            text="Get API Key",
            font=ctk.CTkFont(size=13, underline=True),
            fg_color="transparent",
            text_color=Colors.BLUE_PRIMARY,
            hover_color=Colors.BG_TERTIARY,
            height=30,
            command=self.open_api_key_url
        )
        self.get_key_link.pack(padx=24, anchor="w", pady=(0, 20))

        # Model selection
        model_frame = ctk.CTkFrame(key_card, fg_color="transparent")
        model_frame.pack(fill="x", padx=24, pady=(0, 24))

        model_label = ctk.CTkLabel(
            model_frame,
            text="Select Model:",
            font=ctk.CTkFont(size=14),
            text_color=Colors.TEXT_PRIMARY
        )
        model_label.pack(side="left")

        self.model_dropdown = ctk.CTkComboBox(
            model_frame,
            values=["Gemini 2.5 Pro (Best Quality)"],
            width=280,
            height=40,
            font=ctk.CTkFont(size=13),
            dropdown_font=ctk.CTkFont(size=13),
            border_width=2,
            border_color=Colors.BORDER,
            button_color=Colors.BLUE_PRIMARY,
            button_hover_color=Colors.BLUE_HOVER,
            corner_radius=10
        )
        self.model_dropdown.pack(side="left", padx=(15, 0))

        # Update model dropdown for selected provider
        self.update_model_dropdown("gemini")

        # Status label
        self.setup_status_label = ctk.CTkLabel(
            setup_frame,
            text="",
            font=ctk.CTkFont(size=14),
            text_color=Colors.TEXT_SECONDARY
        )
        self.setup_status_label.pack(pady=(0, 15))

        # Action buttons
        button_frame = ctk.CTkFrame(setup_frame, fg_color="transparent")
        button_frame.pack(fill="x")

        test_btn = ctk.CTkButton(
            button_frame,
            text="Test Connection",
            width=160,
            height=50,
            font=ctk.CTkFont(size=15),
            fg_color=Colors.BG_PRIMARY,
            text_color=Colors.BLUE_PRIMARY,
            border_width=2,
            border_color=Colors.BLUE_PRIMARY,
            hover_color=Colors.BLUE_LIGHT,
            corner_radius=12,
            command=self.test_api_connection
        )
        test_btn.pack(side="left", padx=(0, 15))

        save_btn = ctk.CTkButton(
            button_frame,
            text="Save & Continue",
            width=200,
            height=50,
            font=ctk.CTkFont(size=15, weight="bold"),
            fg_color=Colors.GREEN_PRIMARY,
            hover_color=Colors.GREEN_HOVER,
            corner_radius=12,
            command=self.save_setup
        )
        save_btn.pack(side="left")

        # Select default provider (must be after all widgets are created)
        self.select_provider("gemini")

    def select_provider(self, provider_name: str):
        """Handle provider selection"""
        self.selected_provider.set(provider_name)

        # Update button styles
        for name, btn in self.provider_buttons.items():
            if name == provider_name:
                color = Colors.PROVIDER_COLORS.get(name, Colors.BLUE_PRIMARY)
                btn.configure(
                    fg_color=color,
                    text_color=Colors.BG_PRIMARY,
                    border_color=color
                )
            else:
                btn.configure(
                    fg_color=Colors.BG_PRIMARY,
                    text_color=Colors.TEXT_PRIMARY,
                    border_color=Colors.BORDER
                )

        # Update help text and model dropdown
        info = get_provider_info(provider_name)
        if info:
            self.api_help_label.configure(text=info['api_key_help'])
            self.update_model_dropdown(provider_name)

    def update_model_dropdown(self, provider_name: str):
        """Update the model dropdown for selected provider"""
        info = get_provider_info(provider_name)
        if info:
            models = list(info['models'].values())
            self.model_dropdown.configure(values=models)
            # Set default model
            default_model = info['default_model']
            if default_model in info['models']:
                self.model_dropdown.set(info['models'][default_model])

    def get_selected_model_id(self) -> str:
        """Get the model ID from the display name"""
        provider_name = self.selected_provider.get()
        info = get_provider_info(provider_name)
        if info:
            display_name = self.model_dropdown.get()
            for model_id, name in info['models'].items():
                if name == display_name:
                    return model_id
        return None

    def toggle_key_visibility(self):
        """Toggle API key visibility"""
        if self.api_key_entry.cget("show") == "*":
            self.api_key_entry.configure(show="")
        else:
            self.api_key_entry.configure(show="*")

    def open_api_key_url(self):
        """Open the API key URL in browser"""
        provider_name = self.selected_provider.get()
        info = get_provider_info(provider_name)
        if info:
            webbrowser.open(info['api_key_url'])

    def test_api_connection(self):
        """Test the API connection"""
        provider = self.selected_provider.get()
        api_key = self.api_key_entry.get().strip()
        model = self.get_selected_model_id()

        if not api_key:
            self.setup_status_label.configure(text="Please enter an API key", text_color=Colors.ERROR)
            return

        self.setup_status_label.configure(text="Testing connection...", text_color=Colors.TEXT_SECONDARY)
        self.update()

        def test():
            try:
                provider_instance = get_provider(provider, api_key, model)
                if provider_instance and provider_instance.test_connection():
                    self.after(0, lambda: self.setup_status_label.configure(
                        text="Connection successful!", text_color=Colors.SUCCESS
                    ))
                else:
                    self.after(0, lambda: self.setup_status_label.configure(
                        text="Connection failed. Please check your API key.", text_color=Colors.ERROR
                    ))
            except Exception as e:
                self.after(0, lambda: self.setup_status_label.configure(
                    text=f"Error: {str(e)[:60]}", text_color=Colors.ERROR
                ))

        threading.Thread(target=test, daemon=True).start()

    def save_setup(self):
        """Save the setup configuration"""
        provider = self.selected_provider.get()
        api_key = self.api_key_entry.get().strip()
        model = self.get_selected_model_id()

        if not api_key:
            self.setup_status_label.configure(text="Please enter an API key", text_color=Colors.ERROR)
            return

        try:
            self.api_manager.add_provider(provider, api_key, model, set_as_default=True)
            self.show_main_view()
        except Exception as e:
            self.setup_status_label.configure(text=f"Error saving: {str(e)}", text_color=Colors.ERROR)

    # ===================== MAIN VIEW =====================

    def show_main_view(self):
        """Show the main summarizer view with tab-based navigation"""
        self.clear_container()

        # Header bar with title, provider badge, and theme toggle
        header_bar = ctk.CTkFrame(self.main_container, fg_color=self.theme.get("BG_PRIMARY"), height=60)
        header_bar.pack(fill="x")
        header_bar.pack_propagate(False)

        header_content = ctk.CTkFrame(header_bar, fg_color="transparent")
        header_content.pack(fill="both", expand=True, padx=24)

        # Left side - Logo and title
        title_label = ctk.CTkLabel(
            header_content,
            text="Document Summarizer",
            font=ctk.CTkFont(size=20, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        title_label.pack(side="left", pady=15)

        # Right side - Provider badge and theme toggle
        right_header = ctk.CTkFrame(header_content, fg_color="transparent")
        right_header.pack(side="right", fill="y")

        # Provider badge
        provider_name = self.api_manager.get_default_provider()
        if provider_name:
            info = get_provider_info(provider_name)
            if info:
                provider_color = ThemeManager.PROVIDER_COLORS.get(provider_name, self.theme.get("ACCENT_PRIMARY"))
                provider_badge = ctk.CTkFrame(right_header, fg_color=provider_color, corner_radius=6)
                provider_badge.pack(side="left", padx=(0, 12), pady=14)

                provider_text = ctk.CTkLabel(
                    provider_badge,
                    text=info['display_name'],
                    font=ctk.CTkFont(size=11, weight="bold"),
                    text_color="#FFFFFF"
                )
                provider_text.pack(padx=10, pady=4)

        # Theme toggle button
        theme_icon = "Light" if self.theme.is_dark() else "Dark"
        self.theme_btn = ctk.CTkButton(
            right_header,
            text=theme_icon,
            width=70,
            height=32,
            font=ctk.CTkFont(size=12),
            fg_color=self.theme.get("BG_TERTIARY"),
            text_color=self.theme.get("TEXT_PRIMARY"),
            hover_color=self.theme.get("BORDER"),
            corner_radius=8,
            command=self.toggle_theme
        )
        self.theme_btn.pack(side="left", pady=14)

        # Tab bar
        tab_bar = ctk.CTkFrame(self.main_container, fg_color=self.theme.get("BG_SECONDARY"), height=50)
        tab_bar.pack(fill="x")
        tab_bar.pack_propagate(False)

        tab_container = ctk.CTkFrame(tab_bar, fg_color="transparent")
        tab_container.pack(fill="both", expand=True, padx=24)

        # Tab buttons
        self.tab_buttons = {}
        tabs = [
            ("summarize", "Summarize"),
            ("output", "Output"),
            ("settings", "Settings"),
            ("help", "Help")
        ]

        for tab_id, tab_label in tabs:
            is_active = self.current_tab == tab_id
            btn = ctk.CTkButton(
                tab_container,
                text=tab_label,
                width=100,
                height=36,
                font=ctk.CTkFont(size=13, weight="bold" if is_active else "normal"),
                fg_color=self.theme.get("TAB_ACTIVE") if is_active else "transparent",
                text_color=self.theme.get("ACCENT_PRIMARY") if is_active else self.theme.get("TEXT_SECONDARY"),
                hover_color=self.theme.get("TAB_HOVER"),
                corner_radius=8,
                command=lambda t=tab_id: self.switch_tab(t)
            )
            btn.pack(side="left", padx=(0, 8), pady=7)
            self.tab_buttons[tab_id] = btn

        # Divider
        ctk.CTkFrame(self.main_container, fg_color=self.theme.get("BORDER"), height=1).pack(fill="x")

        # Content area
        self.content_area = ctk.CTkFrame(self.main_container, fg_color=self.theme.get("BG_SECONDARY"))
        self.content_area.pack(fill="both", expand=True)

        # Show current tab content
        self._show_tab_content(self.current_tab)

    def switch_tab(self, tab_id: str):
        """Switch to a different tab"""
        self.current_tab = tab_id

        # Update tab button styles
        for tid, btn in self.tab_buttons.items():
            is_active = tid == tab_id
            btn.configure(
                fg_color=self.theme.get("TAB_ACTIVE") if is_active else "transparent",
                text_color=self.theme.get("ACCENT_PRIMARY") if is_active else self.theme.get("TEXT_SECONDARY"),
                font=ctk.CTkFont(size=13, weight="bold" if is_active else "normal")
            )

        # Clear and show new content
        for widget in self.content_area.winfo_children():
            widget.destroy()
        self._show_tab_content(tab_id)

    def _show_tab_content(self, tab_id: str):
        """Display content for the selected tab"""
        if tab_id == "summarize":
            self._create_summarize_tab()
        elif tab_id == "output":
            self._create_output_tab()
        elif tab_id == "settings":
            self._create_settings_tab()
        elif tab_id == "help":
            self._create_help_tab()

    def toggle_theme(self):
        """Toggle between light and dark themes"""
        self.theme.toggle_theme()
        self.refresh_ui()

    def _create_summarize_tab(self):
        """Create the summarize input tab"""
        # Scrollable content
        scroll_frame = ctk.CTkScrollableFrame(self.content_area, fg_color="transparent")
        scroll_frame.pack(fill="both", expand=True, padx=24, pady=20)

        # Card container
        card = ctk.CTkFrame(scroll_frame, fg_color=self.theme.get("CARD_BG"), corner_radius=12)
        card.pack(fill="x", pady=(0, 16))

        # Upload section header
        header = ctk.CTkFrame(card, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=(20, 12))

        title = ctk.CTkLabel(
            header,
            text="Upload Document",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        title.pack(anchor="w")

        subtitle = ctk.CTkLabel(
            header,
            text="Drag and drop or browse to upload your document",
            font=ctk.CTkFont(size=12),
            text_color=self.theme.get("TEXT_SECONDARY")
        )
        subtitle.pack(anchor="w", pady=(4, 0))

        # Drop zone
        self.drop_zone = ctk.CTkFrame(
            card,
            fg_color=self.theme.get("INPUT_BG"),
            corner_radius=10,
            border_width=2,
            border_color=self.theme.get("BORDER")
        )
        self.drop_zone.pack(fill="x", padx=20, pady=(0, 12))

        # Register drag and drop
        self.drop_zone.drop_target_register(DND_FILES)
        self.drop_zone.dnd_bind('<<Drop>>', self.on_file_drop)

        drop_content = ctk.CTkFrame(self.drop_zone, fg_color="transparent")
        drop_content.pack(expand=True, pady=28)

        # Drop icon
        drop_icon = ctk.CTkLabel(
            drop_content,
            text="[ + ]",
            font=ctk.CTkFont(size=28, weight="bold"),
            text_color=self.theme.get("TEXT_MUTED")
        )
        drop_icon.pack()

        self.drop_label = ctk.CTkLabel(
            drop_content,
            text="Drop your file here",
            font=ctk.CTkFont(size=14),
            text_color=self.theme.get("TEXT_SECONDARY")
        )
        self.drop_label.pack(pady=(8, 12))

        browse_btn = ctk.CTkButton(
            drop_content,
            text="Browse Files",
            width=120,
            height=36,
            font=ctk.CTkFont(size=13),
            fg_color=self.theme.get("ACCENT_PRIMARY"),
            hover_color=self.theme.get("ACCENT_HOVER"),
            corner_radius=8,
            command=self.browse_file
        )
        browse_btn.pack()

        formats_label = ctk.CTkLabel(
            drop_content,
            text="Supports PDF, DOCX, TXT",
            font=ctk.CTkFont(size=11),
            text_color=self.theme.get("TEXT_MUTED")
        )
        formats_label.pack(pady=(10, 0))

        # File info label
        self.file_info_label = ctk.CTkLabel(
            card,
            text="",
            font=ctk.CTkFont(size=12),
            text_color=self.theme.get("SUCCESS")
        )
        self.file_info_label.pack(padx=20, anchor="w")

        # Divider with OR
        or_frame = ctk.CTkFrame(card, fg_color="transparent")
        or_frame.pack(fill="x", padx=20, pady=12)

        or_line1 = ctk.CTkFrame(or_frame, fg_color=self.theme.get("BORDER"), height=1)
        or_line1.pack(side="left", fill="x", expand=True)

        or_label = ctk.CTkLabel(
            or_frame,
            text="OR",
            font=ctk.CTkFont(size=11),
            text_color=self.theme.get("TEXT_MUTED")
        )
        or_label.pack(side="left", padx=12)

        or_line2 = ctk.CTkFrame(or_frame, fg_color=self.theme.get("BORDER"), height=1)
        or_line2.pack(side="left", fill="x", expand=True)

        # Text input section
        text_label = ctk.CTkLabel(
            card,
            text="Paste text directly",
            font=ctk.CTkFont(size=13, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        text_label.pack(padx=20, anchor="w", pady=(0, 8))

        self.text_input = ctk.CTkTextbox(
            card,
            height=100,
            font=ctk.CTkFont(size=12),
            fg_color=self.theme.get("INPUT_BG"),
            text_color=self.theme.get("TEXT_PRIMARY"),
            border_width=1,
            border_color=self.theme.get("BORDER"),
            corner_radius=8
        )
        self.text_input.pack(fill="x", padx=20, pady=(0, 16))

        # Export format selection
        format_frame = ctk.CTkFrame(card, fg_color="transparent")
        format_frame.pack(fill="x", padx=20, pady=(0, 16))

        format_label = ctk.CTkLabel(
            format_frame,
            text="Export Format:",
            font=ctk.CTkFont(size=13),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        format_label.pack(side="left")

        self.format_var = ctk.StringVar(value="pdf")

        pdf_radio = ctk.CTkRadioButton(
            format_frame,
            text="PDF",
            variable=self.format_var,
            value="pdf",
            font=ctk.CTkFont(size=12),
            fg_color=self.theme.get("ACCENT_PRIMARY"),
            hover_color=self.theme.get("ACCENT_HOVER"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        pdf_radio.pack(side="left", padx=(16, 12))

        word_radio = ctk.CTkRadioButton(
            format_frame,
            text="Word",
            variable=self.format_var,
            value="word",
            font=ctk.CTkFont(size=12),
            fg_color=self.theme.get("ACCENT_PRIMARY"),
            hover_color=self.theme.get("ACCENT_HOVER"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        word_radio.pack(side="left")

        # Summarize button
        self.summarize_btn = ctk.CTkButton(
            card,
            text="Summarize Document",
            height=48,
            font=ctk.CTkFont(size=15, weight="bold"),
            fg_color=self.theme.get("SUCCESS"),
            hover_color=self.theme.get("SUCCESS_HOVER"),
            corner_radius=10,
            command=self.summarize_document
        )
        self.summarize_btn.pack(fill="x", padx=20, pady=(0, 20))

    def _create_output_tab(self):
        """Create the output/results tab"""
        scroll_frame = ctk.CTkScrollableFrame(self.content_area, fg_color="transparent")
        scroll_frame.pack(fill="both", expand=True, padx=24, pady=20)

        # Output card
        card = ctk.CTkFrame(scroll_frame, fg_color=self.theme.get("CARD_BG"), corner_radius=12)
        card.pack(fill="x", pady=(0, 16))

        # Header
        header = ctk.CTkFrame(card, fg_color="transparent")
        header.pack(fill="x", padx=20, pady=(20, 12))

        title = ctk.CTkLabel(
            header,
            text="Summary Output",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        title.pack(side="left")

        # Status label
        self.status_label = ctk.CTkLabel(
            card,
            text="No summary generated yet. Go to the Summarize tab to create one.",
            font=ctk.CTkFont(size=12),
            text_color=self.theme.get("TEXT_SECONDARY")
        )
        self.status_label.pack(padx=20, anchor="w", pady=(0, 12))

        # Summary output textbox
        self.summary_output = ctk.CTkTextbox(
            card,
            font=ctk.CTkFont(size=13),
            fg_color=self.theme.get("INPUT_BG"),
            text_color=self.theme.get("TEXT_PRIMARY"),
            border_width=1,
            border_color=self.theme.get("BORDER"),
            corner_radius=8,
            height=280
        )
        self.summary_output.pack(fill="both", expand=True, padx=20, pady=(0, 16))

        # Display existing summary if available
        if self.current_summary:
            self.summary_output.insert("1.0", self.current_summary)
            self.status_label.configure(text="Summary generated!")

        # Spacer
        ctk.CTkFrame(card, fg_color="transparent", height=8).pack()

        # Export Options Card
        export_card = ctk.CTkFrame(scroll_frame, fg_color=self.theme.get("CARD_BG"), corner_radius=12)
        export_card.pack(fill="x")

        export_header = ctk.CTkLabel(
            export_card,
            text="Export Options",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        export_header.pack(pady=(20, 4), padx=20, anchor="w")

        export_subtitle = ctk.CTkLabel(
            export_card,
            text="Customize how your exported document looks",
            font=ctk.CTkFont(size=12),
            text_color=self.theme.get("TEXT_SECONDARY")
        )
        export_subtitle.pack(padx=20, anchor="w", pady=(0, 16))

        # Export Format Selection
        format_section = ctk.CTkFrame(export_card, fg_color="transparent")
        format_section.pack(fill="x", padx=20, pady=(0, 12))

        format_label = ctk.CTkLabel(
            format_section,
            text="File Format:",
            font=ctk.CTkFont(size=13, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        format_label.pack(side="left")

        self.export_format_var = ctk.StringVar(value="pdf")

        pdf_radio = ctk.CTkRadioButton(
            format_section,
            text="PDF",
            variable=self.export_format_var,
            value="pdf",
            font=ctk.CTkFont(size=12),
            fg_color=self.theme.get("ACCENT_PRIMARY"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        pdf_radio.pack(side="left", padx=(16, 12))

        word_radio = ctk.CTkRadioButton(
            format_section,
            text="Word Document",
            variable=self.export_format_var,
            value="word",
            font=ctk.CTkFont(size=12),
            fg_color=self.theme.get("ACCENT_PRIMARY"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        word_radio.pack(side="left")

        # Divider
        ctk.CTkFrame(export_card, fg_color=self.theme.get("BORDER"), height=1).pack(fill="x", padx=20, pady=12)

        # Document Style Selection
        style_label = ctk.CTkLabel(
            export_card,
            text="Document Style:",
            font=ctk.CTkFont(size=13, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        style_label.pack(padx=20, anchor="w", pady=(0, 12))

        self.export_style_var = ctk.StringVar(value="professional")

        # Style options grid
        styles_frame = ctk.CTkFrame(export_card, fg_color="transparent")
        styles_frame.pack(fill="x", padx=20, pady=(0, 16))

        style_options = [
            ("professional", "Professional", "Clean layout with headers and structured paragraphs"),
            ("bullet_points", "Bullet Points", "Key points as a bulleted list for quick reading"),
            ("numbered_list", "Numbered List", "Organized numbered sections for easy reference"),
            ("executive", "Executive Summary", "Concise highlights with bold key takeaways"),
            ("detailed", "Detailed Report", "Comprehensive format with sections and subsections"),
            ("minimalist", "Minimalist", "Simple, clean text with minimal formatting"),
        ]

        for i, (value, label, description) in enumerate(style_options):
            style_option = ctk.CTkFrame(styles_frame, fg_color=self.theme.get("INPUT_BG"), corner_radius=8)
            style_option.pack(fill="x", pady=4)

            radio = ctk.CTkRadioButton(
                style_option,
                text="",
                variable=self.export_style_var,
                value=value,
                width=20,
                fg_color=self.theme.get("ACCENT_PRIMARY"),
                command=lambda v=value: self.export_style_var.set(v)
            )
            radio.pack(side="left", padx=(12, 8), pady=10)

            text_frame = ctk.CTkFrame(style_option, fg_color="transparent")
            text_frame.pack(side="left", fill="x", expand=True, pady=8)

            style_name = ctk.CTkLabel(
                text_frame,
                text=label,
                font=ctk.CTkFont(size=13, weight="bold"),
                text_color=self.theme.get("TEXT_PRIMARY"),
                anchor="w"
            )
            style_name.pack(anchor="w")

            style_desc = ctk.CTkLabel(
                text_frame,
                text=description,
                font=ctk.CTkFont(size=11),
                text_color=self.theme.get("TEXT_MUTED"),
                anchor="w"
            )
            style_desc.pack(anchor="w")

            # Make the whole row clickable
            style_option.bind("<Button-1>", lambda e, v=value: self.export_style_var.set(v))
            text_frame.bind("<Button-1>", lambda e, v=value: self.export_style_var.set(v))
            style_name.bind("<Button-1>", lambda e, v=value: self.export_style_var.set(v))
            style_desc.bind("<Button-1>", lambda e, v=value: self.export_style_var.set(v))

        # Divider
        ctk.CTkFrame(export_card, fg_color=self.theme.get("BORDER"), height=1).pack(fill="x", padx=20, pady=12)

        # Color Scheme Selection
        color_label = ctk.CTkLabel(
            export_card,
            text="Color Scheme:",
            font=ctk.CTkFont(size=13, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        color_label.pack(padx=20, anchor="w", pady=(0, 12))

        self.export_color_var = ctk.StringVar(value="blue")

        colors_frame = ctk.CTkFrame(export_card, fg_color="transparent")
        colors_frame.pack(fill="x", padx=20, pady=(0, 16))

        color_options = [
            ("blue", "Blue", "#2563EB"),
            ("green", "Green", "#059669"),
            ("purple", "Purple", "#7C3AED"),
            ("red", "Red", "#DC2626"),
            ("orange", "Orange", "#EA580C"),
            ("gray", "Grayscale", "#475569"),
        ]

        for value, label, color in color_options:
            color_btn = ctk.CTkButton(
                colors_frame,
                text=label,
                width=80,
                height=36,
                font=ctk.CTkFont(size=11),
                fg_color=color if self.export_color_var.get() == value else self.theme.get("INPUT_BG"),
                text_color="#FFFFFF" if self.export_color_var.get() == value else self.theme.get("TEXT_PRIMARY"),
                hover_color=color,
                corner_radius=8,
                command=lambda v=value: self._select_export_color(v)
            )
            color_btn.pack(side="left", padx=(0, 8))
            # Store reference for updating
            if not hasattr(self, 'color_buttons'):
                self.color_buttons = {}
            self.color_buttons[value] = (color_btn, color)

        # Export button
        self.export_btn = ctk.CTkButton(
            export_card,
            text="Export Summary",
            height=48,
            font=ctk.CTkFont(size=15, weight="bold"),
            fg_color=self.theme.get("SUCCESS"),
            hover_color=self.theme.get("SUCCESS_HOVER"),
            corner_radius=10,
            state="normal" if self.current_summary else "disabled",
            command=self.export_summary
        )
        self.export_btn.pack(fill="x", padx=20, pady=(8, 20))

    def _select_export_color(self, color_value: str):
        """Handle color scheme selection"""
        self.export_color_var.set(color_value)
        # Update button styles
        for value, (btn, color) in self.color_buttons.items():
            if value == color_value:
                btn.configure(fg_color=color, text_color="#FFFFFF")
            else:
                btn.configure(fg_color=self.theme.get("INPUT_BG"), text_color=self.theme.get("TEXT_PRIMARY"))

    def _create_settings_tab(self):
        """Create the settings tab"""
        scroll_frame = ctk.CTkScrollableFrame(self.content_area, fg_color="transparent")
        scroll_frame.pack(fill="both", expand=True, padx=24, pady=20)

        # Configured providers card
        providers_card = ctk.CTkFrame(scroll_frame, fg_color=self.theme.get("CARD_BG"), corner_radius=12)
        providers_card.pack(fill="x", pady=(0, 16))

        providers_header = ctk.CTkLabel(
            providers_card,
            text="Configured AI Providers",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        providers_header.pack(pady=(20, 12), padx=20, anchor="w")

        providers = self.api_manager.list_providers()

        if providers:
            for prov in providers:
                prov_frame = ctk.CTkFrame(providers_card, fg_color=self.theme.get("INPUT_BG"), corner_radius=8)
                prov_frame.pack(fill="x", padx=20, pady=4)

                info = get_provider_info(prov['name'])
                color = ThemeManager.PROVIDER_COLORS.get(prov['name'], self.theme.get("ACCENT_PRIMARY"))

                # Provider color indicator
                color_bar = ctk.CTkFrame(prov_frame, fg_color=color, width=4, corner_radius=2)
                color_bar.pack(side="left", fill="y", padx=(0, 12), pady=8)

                details = ctk.CTkFrame(prov_frame, fg_color="transparent")
                details.pack(side="left", fill="x", expand=True, pady=10)

                name_row = ctk.CTkFrame(details, fg_color="transparent")
                name_row.pack(fill="x")

                name_label = ctk.CTkLabel(
                    name_row,
                    text=info['display_name'] if info else prov['name'],
                    font=ctk.CTkFont(size=14, weight="bold"),
                    text_color=self.theme.get("TEXT_PRIMARY")
                )
                name_label.pack(side="left")

                if prov['is_default']:
                    default_badge = ctk.CTkLabel(
                        name_row,
                        text="DEFAULT",
                        font=ctk.CTkFont(size=9, weight="bold"),
                        text_color="#FFFFFF",
                        fg_color=self.theme.get("SUCCESS"),
                        corner_radius=4
                    )
                    default_badge.pack(side="left", padx=8)

                model_name = "Default model"
                if prov['model'] and info:
                    model_name = info['models'].get(prov['model'], prov['model'])

                model_label = ctk.CTkLabel(
                    details,
                    text=f"Model: {model_name}  |  Key: {prov['api_key_preview']}",
                    font=ctk.CTkFont(size=11),
                    text_color=self.theme.get("TEXT_SECONDARY")
                )
                model_label.pack(anchor="w")

                # Action buttons
                actions = ctk.CTkFrame(prov_frame, fg_color="transparent")
                actions.pack(side="right", padx=12, pady=10)

                if not prov['is_default']:
                    set_default_btn = ctk.CTkButton(
                        actions,
                        text="Set Default",
                        width=80,
                        height=28,
                        font=ctk.CTkFont(size=11),
                        fg_color=self.theme.get("ACCENT_PRIMARY"),
                        hover_color=self.theme.get("ACCENT_HOVER"),
                        corner_radius=6,
                        command=lambda p=prov['name']: self._set_default(p)
                    )
                    set_default_btn.pack(side="left", padx=(0, 6))

                remove_btn = ctk.CTkButton(
                    actions,
                    text="Remove",
                    width=70,
                    height=28,
                    font=ctk.CTkFont(size=11),
                    fg_color=self.theme.get("ERROR"),
                    hover_color="#DC2626",
                    corner_radius=6,
                    command=lambda p=prov['name']: self._remove_provider(p)
                )
                remove_btn.pack(side="left")
        else:
            no_providers = ctk.CTkLabel(
                providers_card,
                text="No providers configured. Add one below.",
                font=ctk.CTkFont(size=13),
                text_color=self.theme.get("TEXT_SECONDARY")
            )
            no_providers.pack(pady=16)

        ctk.CTkFrame(providers_card, fg_color="transparent", height=16).pack()

        # Add new provider card
        add_card = ctk.CTkFrame(scroll_frame, fg_color=self.theme.get("CARD_BG"), corner_radius=12)
        add_card.pack(fill="x", pady=(0, 16))

        add_header = ctk.CTkLabel(
            add_card,
            text="Add New Provider",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        add_header.pack(pady=(20, 12), padx=20, anchor="w")

        # Provider selection
        provider_select_frame = ctk.CTkFrame(add_card, fg_color="transparent")
        provider_select_frame.pack(fill="x", padx=20, pady=(0, 12))

        self.new_provider_var = ctk.StringVar(value="gemini")

        all_providers = get_all_providers()
        col = 0
        for name, info in all_providers.items():
            rb = ctk.CTkRadioButton(
                provider_select_frame,
                text=info['display_name'],
                variable=self.new_provider_var,
                value=name,
                font=ctk.CTkFont(size=12),
                fg_color=self.theme.get("ACCENT_PRIMARY"),
                text_color=self.theme.get("TEXT_PRIMARY")
            )
            rb.grid(row=0, column=col, padx=8, pady=4, sticky="w")
            col += 1

        # API key input
        input_frame = ctk.CTkFrame(add_card, fg_color="transparent")
        input_frame.pack(fill="x", padx=20, pady=(0, 20))

        self.new_api_key_entry = ctk.CTkEntry(
            input_frame,
            placeholder_text="Enter API key...",
            width=320,
            height=40,
            font=ctk.CTkFont(size=12),
            show="*",
            border_width=1,
            border_color=self.theme.get("BORDER"),
            fg_color=self.theme.get("INPUT_BG"),
            text_color=self.theme.get("TEXT_PRIMARY"),
            corner_radius=8
        )
        self.new_api_key_entry.pack(side="left", padx=(0, 10))

        add_btn = ctk.CTkButton(
            input_frame,
            text="Add Provider",
            width=110,
            height=40,
            font=ctk.CTkFont(size=13),
            fg_color=self.theme.get("SUCCESS"),
            hover_color=self.theme.get("SUCCESS_HOVER"),
            corner_radius=8,
            command=self._add_provider
        )
        add_btn.pack(side="left")

        # Theme settings card
        theme_card = ctk.CTkFrame(scroll_frame, fg_color=self.theme.get("CARD_BG"), corner_radius=12)
        theme_card.pack(fill="x")

        theme_header = ctk.CTkLabel(
            theme_card,
            text="Appearance",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        theme_header.pack(pady=(20, 12), padx=20, anchor="w")

        theme_row = ctk.CTkFrame(theme_card, fg_color="transparent")
        theme_row.pack(fill="x", padx=20, pady=(0, 20))

        theme_label = ctk.CTkLabel(
            theme_row,
            text="Theme:",
            font=ctk.CTkFont(size=13),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        theme_label.pack(side="left")

        current_theme_text = "Dark Mode" if self.theme.is_dark() else "Light Mode"
        theme_toggle_btn = ctk.CTkButton(
            theme_row,
            text=f"Switch to {'Light' if self.theme.is_dark() else 'Dark'} Mode",
            width=160,
            height=36,
            font=ctk.CTkFont(size=12),
            fg_color=self.theme.get("BG_TERTIARY"),
            text_color=self.theme.get("TEXT_PRIMARY"),
            hover_color=self.theme.get("BORDER"),
            corner_radius=8,
            command=self.toggle_theme
        )
        theme_toggle_btn.pack(side="left", padx=(12, 0))

        current_label = ctk.CTkLabel(
            theme_row,
            text=f"Currently: {current_theme_text}",
            font=ctk.CTkFont(size=11),
            text_color=self.theme.get("TEXT_MUTED")
        )
        current_label.pack(side="left", padx=(12, 0))

    def _create_help_tab(self):
        """Create the help tab with API key instructions"""
        scroll_frame = ctk.CTkScrollableFrame(self.content_area, fg_color="transparent")
        scroll_frame.pack(fill="both", expand=True, padx=24, pady=20)

        # Intro card
        intro_card = ctk.CTkFrame(scroll_frame, fg_color=self.theme.get("CARD_BG"), corner_radius=12)
        intro_card.pack(fill="x", pady=(0, 16))

        intro_header = ctk.CTkLabel(
            intro_card,
            text="Getting Started",
            font=ctk.CTkFont(size=16, weight="bold"),
            text_color=self.theme.get("TEXT_PRIMARY")
        )
        intro_header.pack(pady=(20, 8), padx=20, anchor="w")

        intro_text = ctk.CTkLabel(
            intro_card,
            text="To use Document Summarizer, you need an API key from one of the supported AI providers.\nBelow are instructions for getting a free or paid API key from each provider.",
            font=ctk.CTkFont(size=12),
            text_color=self.theme.get("TEXT_SECONDARY"),
            justify="left"
        )
        intro_text.pack(padx=20, pady=(0, 20), anchor="w")

        # Provider help cards
        providers_info = [
            {
                "name": "Google Gemini",
                "color": ThemeManager.PROVIDER_COLORS["gemini"],
                "free": True,
                "url": "https://aistudio.google.com/app/apikey",
                "steps": [
                    "1. Go to Google AI Studio (aistudio.google.com)",
                    "2. Sign in with your Google account",
                    "3. Click 'Get API Key' in the sidebar",
                    "4. Create a new API key",
                    "5. Copy and paste it here"
                ],
                "notes": "Free tier: 15 requests/minute. Recommended for most users!"
            },
            {
                "name": "OpenAI (GPT)",
                "color": ThemeManager.PROVIDER_COLORS["openai"],
                "free": False,
                "url": "https://platform.openai.com/api-keys",
                "steps": [
                    "1. Go to platform.openai.com",
                    "2. Sign up or log in",
                    "3. Navigate to API Keys",
                    "4. Create a new secret key",
                    "5. Add a payment method"
                ],
                "notes": "Paid only. Requires billing setup."
            },
            {
                "name": "Anthropic Claude",
                "color": ThemeManager.PROVIDER_COLORS["claude"],
                "free": False,
                "url": "https://console.anthropic.com/settings/keys",
                "steps": [
                    "1. Go to console.anthropic.com",
                    "2. Sign up and verify email",
                    "3. Navigate to API Keys",
                    "4. Create a new key",
                    "5. Add credits in Billing"
                ],
                "notes": "Paid only. High-quality responses."
            },
            {
                "name": "Groq",
                "color": ThemeManager.PROVIDER_COLORS["groq"],
                "free": True,
                "url": "https://console.groq.com/keys",
                "steps": [
                    "1. Go to console.groq.com",
                    "2. Sign up for free",
                    "3. Navigate to API Keys",
                    "4. Create a new key",
                    "5. No payment needed!"
                ],
                "notes": "Free tier available! Very fast."
            }
        ]

        for prov in providers_info:
            card = ctk.CTkFrame(scroll_frame, fg_color=self.theme.get("CARD_BG"), corner_radius=12)
            card.pack(fill="x", pady=(0, 12))

            # Header
            card_header = ctk.CTkFrame(card, fg_color="transparent")
            card_header.pack(fill="x", padx=20, pady=(16, 8))

            color_bar = ctk.CTkFrame(card_header, fg_color=prov["color"], width=4, height=24, corner_radius=2)
            color_bar.pack(side="left", padx=(0, 12))

            name_label = ctk.CTkLabel(
                card_header,
                text=prov["name"],
                font=ctk.CTkFont(size=15, weight="bold"),
                text_color=self.theme.get("TEXT_PRIMARY")
            )
            name_label.pack(side="left")

            # Free/Paid badge
            badge_text = "FREE" if prov["free"] else "PAID"
            badge_color = self.theme.get("SUCCESS") if prov["free"] else self.theme.get("WARNING")
            badge = ctk.CTkLabel(
                card_header,
                text=badge_text,
                font=ctk.CTkFont(size=10, weight="bold"),
                text_color="#FFFFFF",
                fg_color=badge_color,
                corner_radius=4
            )
            badge.pack(side="left", padx=10)

            # Get API Key button
            get_key_btn = ctk.CTkButton(
                card_header,
                text="Get API Key",
                width=90,
                height=28,
                font=ctk.CTkFont(size=11),
                fg_color=prov["color"],
                hover_color=self.theme.get("ACCENT_HOVER"),
                corner_radius=6,
                command=lambda url=prov["url"]: webbrowser.open(url)
            )
            get_key_btn.pack(side="right")

            # Steps
            steps_frame = ctk.CTkFrame(card, fg_color=self.theme.get("INPUT_BG"), corner_radius=8)
            steps_frame.pack(fill="x", padx=20, pady=(0, 8))

            for step in prov["steps"]:
                step_label = ctk.CTkLabel(
                    steps_frame,
                    text=step,
                    font=ctk.CTkFont(size=11),
                    text_color=self.theme.get("TEXT_PRIMARY"),
                    anchor="w"
                )
                step_label.pack(anchor="w", padx=12, pady=2)

            # Notes
            notes_label = ctk.CTkLabel(
                card,
                text=prov["notes"],
                font=ctk.CTkFont(size=11),
                text_color=self.theme.get("TEXT_MUTED")
            )
            notes_label.pack(padx=20, pady=(0, 16), anchor="w")

    def on_file_drop(self, event):
        """Handle file drop"""
        file_path = event.data.strip('{}')
        self.load_file(file_path)

    def browse_file(self):
        """Open file browser"""
        file_path = filedialog.askopenfilename(
            title="Select Document",
            filetypes=[
                ("All Supported", "*.pdf *.docx *.doc *.txt"),
                ("PDF Files", "*.pdf"),
                ("Word Documents", "*.docx *.doc"),
                ("Text Files", "*.txt")
            ]
        )
        if file_path:
            self.load_file(file_path)

    def load_file(self, file_path: str):
        """Load and extract text from file"""
        if not os.path.exists(file_path):
            messagebox.showerror("Error", "File not found")
            return

        ext = file_path.rsplit('.', 1)[-1].lower()
        if ext not in ['pdf', 'docx', 'doc', 'txt']:
            messagebox.showerror("Error", "Unsupported format. Use PDF, DOCX, or TXT.")
            return

        self.current_file_path = file_path
        filename = os.path.basename(file_path)
        self.original_filename = filename
        self.file_info_label.configure(text=f"Loaded: {filename}")
        self.status_label.configure(text="Extracting text...")
        self.update()

        def extract():
            try:
                text = self.extract_text(file_path, filename)
                self.extracted_text = text
                char_count = len(text)
                self.after(0, lambda: self.status_label.configure(
                    text=f"Ready to summarize ({char_count:,} characters)"
                ))
            except Exception as e:
                self.after(0, lambda: self.status_label.configure(text=f"Error: {str(e)}"))

        threading.Thread(target=extract, daemon=True).start()

    def extract_text(self, file_path: str, filename: str) -> str:
        """Extract text from document"""
        ext = filename.rsplit('.', 1)[-1].lower()

        if ext == 'pdf':
            return self._extract_pdf(file_path)
        elif ext in ['docx', 'doc']:
            return self._extract_docx(file_path)
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        raise Exception(f"Unsupported: {ext}")

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        return text.strip()

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = Document(file_path)
        return '\n'.join([p.text for p in doc.paragraphs])

    def summarize_document(self):
        """Summarize the document"""
        text = self.extracted_text or self.text_input.get("1.0", "end-1c").strip()

        if not text or len(text) < 50:
            messagebox.showwarning("Warning", "Please provide more text (at least 50 characters).")
            return

        provider_name = self.api_manager.get_default_provider()
        if not provider_name:
            messagebox.showerror("Error", "No AI provider configured.")
            return

        api_key = self.api_manager.get_api_key(provider_name)
        model = self.api_manager.get_model(provider_name)

        self.summarize_btn.configure(state="disabled", text="Summarizing...")
        self.update()

        def do_summarize():
            try:
                provider = get_provider(provider_name, api_key, model)
                summary = provider.summarize(text)
                self.current_summary = summary

                # Switch to output tab and show summary
                def show_result():
                    self.summarize_btn.configure(state="normal", text="Summarize Document")
                    self.switch_tab("output")
                    messagebox.showinfo("Success", "Summary generated! View it in the Output tab.")

                self.after(0, show_result)
            except Exception as e:
                def show_error():
                    self.summarize_btn.configure(state="normal", text="Summarize Document")
                    messagebox.showerror("Error", f"Failed to generate summary: {str(e)}")

                self.after(0, show_error)

        threading.Thread(target=do_summarize, daemon=True).start()

    def export_summary(self):
        """Export summary to file"""
        if not self.current_summary:
            return

        # Get export options
        output_format = getattr(self, 'export_format_var', ctk.StringVar(value="pdf")).get()
        style = getattr(self, 'export_style_var', ctk.StringVar(value="professional")).get()
        color = getattr(self, 'export_color_var', ctk.StringVar(value="blue")).get()

        ext = ".pdf" if output_format == "pdf" else ".docx"
        base = self.original_filename.rsplit('.', 1)[0]

        file_path = filedialog.asksaveasfilename(
            title="Save Summary",
            defaultextension=ext,
            initialfile=f"summary_{base}",
            filetypes=[("PDF Files", "*.pdf") if output_format == "pdf" else ("Word Documents", "*.docx")]
        )

        if not file_path:
            return

        try:
            if output_format == "pdf":
                self._create_pdf(self.current_summary, file_path, style, color)
            else:
                self._create_docx(self.current_summary, file_path, style, color)
            messagebox.showinfo("Success", f"Saved to:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def _get_color_rgb(self, color_name: str):
        """Get RGB color values for the selected color scheme"""
        from reportlab.lib.colors import HexColor
        colors = {
            "blue": {"primary": "#2563EB", "secondary": "#DBEAFE", "dark": "#1E40AF"},
            "green": {"primary": "#059669", "secondary": "#D1FAE5", "dark": "#047857"},
            "purple": {"primary": "#7C3AED", "secondary": "#EDE9FE", "dark": "#6D28D9"},
            "red": {"primary": "#DC2626", "secondary": "#FEE2E2", "dark": "#B91C1C"},
            "orange": {"primary": "#EA580C", "secondary": "#FFEDD5", "dark": "#C2410C"},
            "gray": {"primary": "#475569", "secondary": "#F1F5F9", "dark": "#334155"},
        }
        return colors.get(color_name, colors["blue"])

    def _create_pdf(self, summary: str, path: str, style: str = "professional", color: str = "blue"):
        """Create styled PDF"""
        from reportlab.lib.colors import HexColor, black, white
        from reportlab.lib.units import inch
        from reportlab.platypus import ListFlowable, ListItem, HRFlowable, Table, TableStyle

        color_scheme = self._get_color_rgb(color)
        primary_color = HexColor(color_scheme["primary"])
        secondary_color = HexColor(color_scheme["secondary"])
        dark_color = HexColor(color_scheme["dark"])

        doc = SimpleDocTemplate(
            path,
            pagesize=letter,
            rightMargin=60,
            leftMargin=60,
            topMargin=50,
            bottomMargin=40
        )

        styles = getSampleStyleSheet()

        # Custom styles based on color scheme
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=primary_color,
            spaceAfter=20,
            fontName='Helvetica-Bold'
        ))

        styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=dark_color,
            spaceBefore=16,
            spaceAfter=8,
            fontName='Helvetica-Bold'
        ))

        styles.add(ParagraphStyle(
            name='CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            textColor=black,
            spaceAfter=8,
            leading=16
        ))

        styles.add(ParagraphStyle(
            name='BulletText',
            parent=styles['BodyText'],
            fontSize=11,
            textColor=black,
            leftIndent=20,
            spaceAfter=6,
            bulletIndent=10
        ))

        styles.add(ParagraphStyle(
            name='KeyPoint',
            parent=styles['BodyText'],
            fontSize=11,
            textColor=black,
            backColor=secondary_color,
            borderPadding=8,
            spaceAfter=10,
            leading=16
        ))

        elements = []

        # Title
        elements.append(Paragraph(f"Summary: {self.original_filename}", styles['CustomTitle']))

        # Colored line under title
        elements.append(HRFlowable(
            width="100%",
            thickness=3,
            color=primary_color,
            spaceAfter=20
        ))

        # Process content based on style
        lines = [line.strip() for line in summary.split('\n') if line.strip()]

        if style == "professional":
            for line in lines:
                if line.startswith('#') or line.startswith('**') and line.endswith('**'):
                    clean_line = line.replace('#', '').replace('**', '').strip()
                    elements.append(Paragraph(clean_line, styles['CustomHeading']))
                else:
                    elements.append(Paragraph(line, styles['CustomBody']))

        elif style == "bullet_points":
            elements.append(Paragraph("Key Points", styles['CustomHeading']))
            bullet_items = []
            for line in lines:
                clean_line = line.lstrip('-•* ').replace('**', '').replace('#', '').strip()
                if clean_line:
                    bullet_items.append(ListItem(Paragraph(clean_line, styles['BulletText'])))
            if bullet_items:
                elements.append(ListFlowable(bullet_items, bulletType='bullet', bulletColor=primary_color))

        elif style == "numbered_list":
            elements.append(Paragraph("Summary Points", styles['CustomHeading']))
            numbered_items = []
            for i, line in enumerate(lines, 1):
                clean_line = line.lstrip('-•*0123456789. ').replace('**', '').replace('#', '').strip()
                if clean_line:
                    numbered_items.append(ListItem(Paragraph(clean_line, styles['BulletText'])))
            if numbered_items:
                elements.append(ListFlowable(numbered_items, bulletType='1', bulletColor=primary_color))

        elif style == "executive":
            elements.append(Paragraph("Executive Summary", styles['CustomHeading']))
            elements.append(Spacer(1, 10))

            # First few lines as key highlights
            key_points = lines[:min(5, len(lines))]
            for point in key_points:
                clean_point = point.lstrip('-•* ').replace('**', '').replace('#', '').strip()
                if clean_point:
                    elements.append(Paragraph(f"<b>→</b> {clean_point}", styles['KeyPoint']))

            if len(lines) > 5:
                elements.append(Spacer(1, 15))
                elements.append(Paragraph("Additional Details", styles['CustomHeading']))
                for line in lines[5:]:
                    clean_line = line.replace('**', '').replace('#', '').strip()
                    if clean_line:
                        elements.append(Paragraph(clean_line, styles['CustomBody']))

        elif style == "detailed":
            section_num = 1
            for line in lines:
                if line.startswith('#') or (line.startswith('**') and line.endswith('**')):
                    clean_line = line.replace('#', '').replace('**', '').strip()
                    elements.append(Paragraph(f"Section {section_num}: {clean_line}", styles['CustomHeading']))
                    section_num += 1
                elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
                    clean_line = line.lstrip('-•* ').strip()
                    elements.append(Paragraph(f"    • {clean_line}", styles['CustomBody']))
                else:
                    elements.append(Paragraph(line, styles['CustomBody']))

        elif style == "minimalist":
            for line in lines:
                clean_line = line.replace('**', '').replace('#', '').replace('-', '').replace('•', '').strip()
                if clean_line:
                    elements.append(Paragraph(clean_line, styles['CustomBody']))
                    elements.append(Spacer(1, 6))

        # Footer
        elements.append(Spacer(1, 30))
        elements.append(HRFlowable(width="100%", thickness=1, color=HexColor("#E2E8F0")))
        elements.append(Spacer(1, 10))

        footer_style = ParagraphStyle(
            name='Footer',
            parent=styles['Normal'],
            fontSize=9,
            textColor=HexColor("#94A3B8"),
            alignment=1  # Center
        )
        elements.append(Paragraph("Generated by Document Summarizer", footer_style))

        doc.build(elements)

    def _create_docx(self, summary: str, path: str, style: str = "professional", color: str = "blue"):
        """Create styled DOCX"""
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        color_map = {
            "blue": RGBColor(37, 99, 235),
            "green": RGBColor(5, 150, 105),
            "purple": RGBColor(124, 58, 237),
            "red": RGBColor(220, 38, 38),
            "orange": RGBColor(234, 88, 12),
            "gray": RGBColor(71, 85, 105),
        }
        accent_color = color_map.get(color, color_map["blue"])

        doc = Document()

        # Title
        title = doc.add_heading(f"Summary: {self.original_filename}", 0)
        title.runs[0].font.color.rgb = accent_color

        lines = [line.strip() for line in summary.split('\n') if line.strip()]

        if style == "professional":
            for line in lines:
                if line.startswith('#') or (line.startswith('**') and line.endswith('**')):
                    clean_line = line.replace('#', '').replace('**', '').strip()
                    heading = doc.add_heading(clean_line, level=2)
                    heading.runs[0].font.color.rgb = accent_color
                else:
                    doc.add_paragraph(line)

        elif style == "bullet_points":
            doc.add_heading("Key Points", level=1).runs[0].font.color.rgb = accent_color
            for line in lines:
                clean_line = line.lstrip('-•* ').replace('**', '').replace('#', '').strip()
                if clean_line:
                    para = doc.add_paragraph(clean_line, style='List Bullet')

        elif style == "numbered_list":
            doc.add_heading("Summary Points", level=1).runs[0].font.color.rgb = accent_color
            for line in lines:
                clean_line = line.lstrip('-•*0123456789. ').replace('**', '').replace('#', '').strip()
                if clean_line:
                    doc.add_paragraph(clean_line, style='List Number')

        elif style == "executive":
            doc.add_heading("Executive Summary", level=1).runs[0].font.color.rgb = accent_color

            key_points = lines[:min(5, len(lines))]
            for point in key_points:
                clean_point = point.lstrip('-•* ').replace('**', '').replace('#', '').strip()
                if clean_point:
                    para = doc.add_paragraph()
                    run = para.add_run(f"→ {clean_point}")
                    run.bold = True

            if len(lines) > 5:
                doc.add_heading("Additional Details", level=2).runs[0].font.color.rgb = accent_color
                for line in lines[5:]:
                    clean_line = line.replace('**', '').replace('#', '').strip()
                    if clean_line:
                        doc.add_paragraph(clean_line)

        elif style == "detailed":
            section_num = 1
            for line in lines:
                if line.startswith('#') or (line.startswith('**') and line.endswith('**')):
                    clean_line = line.replace('#', '').replace('**', '').strip()
                    heading = doc.add_heading(f"Section {section_num}: {clean_line}", level=2)
                    heading.runs[0].font.color.rgb = accent_color
                    section_num += 1
                elif line.startswith('-') or line.startswith('•') or line.startswith('*'):
                    clean_line = line.lstrip('-•* ').strip()
                    doc.add_paragraph(clean_line, style='List Bullet')
                else:
                    doc.add_paragraph(line)

        elif style == "minimalist":
            for line in lines:
                clean_line = line.replace('**', '').replace('#', '').replace('-', '').replace('•', '').strip()
                if clean_line:
                    para = doc.add_paragraph(clean_line)
                    para.paragraph_format.space_after = Pt(12)

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph("Generated by Document Summarizer")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(148, 163, 184)

        doc.save(path)

    def _set_default(self, provider_name: str):
        """Set provider as default"""
        self.api_manager.set_default_provider(provider_name)
        self.switch_tab("settings")

    def _remove_provider(self, provider_name: str):
        """Remove a provider"""
        if messagebox.askyesno("Confirm", f"Remove {provider_name}?"):
            self.api_manager.remove_provider(provider_name)
            if not self.api_manager.has_any_provider():
                self.show_setup_view()
            else:
                self.switch_tab("settings")

    def _add_provider(self):
        """Add a new provider"""
        provider = self.new_provider_var.get()
        api_key = self.new_api_key_entry.get().strip()

        if not api_key:
            messagebox.showwarning("Warning", "Please enter an API key")
            return

        info = get_provider_info(provider)
        default_model = info['default_model'] if info else None

        try:
            self.api_manager.add_provider(provider, api_key, default_model, set_as_default=False)
            self.switch_tab("settings")
        except Exception as e:
            messagebox.showerror("Error", str(e))


def main():
    """Main entry point"""
    app = DocumentSummarizerApp()
    app.mainloop()


if __name__ == '__main__':
    main()
